from flask import Flask, render_template, request, redirect, url_for, send_file, flash, session
import os
import json
import pandas as pd
from datetime import date, datetime
from docx import Document
import locale
from werkzeug.utils import secure_filename
import re
import os
import json
from werkzeug.security import generate_password_hash, check_password_hash
from flask_mail import Mail, Message
from apscheduler.schedulers.background import BackgroundScheduler

app = Flask(__name__)
app.secret_key = 'clave_secreta'

# Archivo donde se almacenan las actualizaciones del archivo
LOG_FILE = "actualizaciones.json"
# Archivo donde se almacenan los usuarios
USERS_FILE = 'usuarios.json'
# Archivo donde se almacenan los eventos del calendario
EVENTOS_FILE = 'eventos.json'

app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
app.config['TRANSACCIONES_FOLDER'] = 'transacciones'
os.makedirs(app.config['TRANSACCIONES_FOLDER'], exist_ok=True)

app.config['MAIL_SERVER'] = 'smtp.gmail.com'  
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'anacevazeditor@gmail.com'
app.config['MAIL_PASSWORD'] = 'nwyw ukae bkhq uush'  

mail = Mail(app)

# Asegurarse de que exista el archivo de usuarios
if not os.path.exists(USERS_FILE):
    with open(USERS_FILE, 'w') as f:
        json.dump([], f)

# Asegurarse de que exista el archivo de eventos
if not os.path.exists(EVENTOS_FILE):
    with open(EVENTOS_FILE, 'w') as f:
        json.dump([], f)


FECHA_RANGO_SIMPLE = re.compile(
    r"(?i)(?:del\s+)?(\d{1,2})\s*(?:al|–|-)\s*(\d{1,2})\s*(?:de)?\s*([a-zA-ZñÑ]+)\s*(?:de[l]*?|,)?\s*(\d{4})"
)

FECHA_RANGO_DOBLE_MES = re.compile(
    r"(?i)(\d{1,2})\s*(?:de)?\s*([a-zA-ZñÑ]+)\s*(?:al|–|-)\s*(\d{1,2})\s*(?:de)?\s*([a-zA-ZñÑ]+)\s*(?:de[l]*?|,)?\s*(\d{4})"
)

FECHA_UNICA = re.compile(
    r"(?i)(\d{1,2})\s*(?:de)?\s*([a-zA-ZñÑ]+)\s*(?:de[l]*?|,)?\s*(\d{4})"
)

PRECIO_REGEX = re.compile(
    r"\$([0-9]+(?:\.[0-9]{1,2})?)"
)
IMPUESTO_REGEX = re.compile(
    r"(\d{1,2}(?:\.\d{1,2})?)\s*%.*?(?:impuesto[s]?|IGTF)", re.IGNORECASE
)
PRECIO_TOTAL_REGEX = re.compile(
    r"monto a pagar (?:sería|es|seria|seran|serán)[^\$]*\$([0-9]+(?:\.[0-9]{1,2})?)", re.IGNORECASE
)

MESES = {
    "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
    "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
    "septiembre": "09", "setiembre": "09", "octubre": "10",
    "noviembre": "11", "diciembre": "12"
}


def registrar_actualizacion():
    hoy = datetime.today().strftime('%Y-%m-%d')
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, 'r') as f:
            fechas = json.load(f)
    else:
        fechas = []

    fechas.append(hoy)

    with open(LOG_FILE, 'w') as f:
        json.dump(fechas, f)
        
def parse_fecha(dia, mes_nombre, año):
    mes = MESES.get(mes_nombre.lower())
    if mes:
        return f"{año}-{mes}-{int(dia):02d}"
    return None

def cargar_usuarios():
    with open(USERS_FILE, 'r') as f:
        return json.load(f)

def guardar_usuarios(usuarios):
    with open(USERS_FILE, 'w') as f:
        json.dump(usuarios, f, indent=4)

def buscar_usuario_por_email(email):
    usuarios = cargar_usuarios()
    for usuario in usuarios:
        if usuario['email'] == email:
            return usuario
    return None


@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        usuario = buscar_usuario_por_email(email)

        if usuario and check_password_hash(usuario['password'], password):
            session['usuario'] = usuario['email']
            flash('Inicio de sesión exitoso.', 'success')
            return redirect(url_for('archivo'))
        else:
            flash('Credenciales inválidas.', 'error')
            return redirect(url_for('login'))

    return render_template('login.html')

@app.route('/sign-up', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        email = request.form['email']
        first_name = request.form['first_name']
        last_name = request.form['last_name']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        pregunta_recuperacion = request.form['pregunta_recuperacion']
        respuesta_recuperacion = request.form['respuesta_recuperacion']

        if password != confirm_password:
            flash('Las contraseñas no coinciden.', 'error')
            return redirect(url_for('register'))

        if buscar_usuario_por_email(email):
            flash('El usuario ya está registrado.', 'error')
            return redirect(url_for('register'))

        usuarios = cargar_usuarios()
        hashed_password = generate_password_hash(password)

        usuarios.append({
            'email': email,
            'first_name': first_name,
            'last_name': last_name,
            'password': hashed_password,
            'pregunta_recuperacion': pregunta_recuperacion,
            'respuesta_recuperacion': respuesta_recuperacion
        })
        guardar_usuarios(usuarios)

        flash('Registro exitoso. Por favor inicia sesión.', 'success')
        return redirect(url_for('login'))

    return render_template('sign-up.html')

@app.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        usuario = buscar_usuario_por_email(email)
        if usuario:
            session['recovery_email'] = email
            return redirect(url_for('security_question'))
        else:
            flash('Correo no registrado.', 'error')
    return render_template('forgot-password.html')


@app.route('/security-question', methods=['GET', 'POST'])
def security_question():
    email = session.get('recovery_email')
    if not email:
        return redirect(url_for('forgot_password'))

    usuario = buscar_usuario_por_email(email)
    pregunta = usuario['pregunta_recuperacion']

    if request.method == 'POST':
        respuesta = request.form['respuesta']
        if respuesta.strip().lower() == usuario['respuesta_recuperacion'].strip().lower():
            session['reset_allowed'] = True
            return redirect(url_for('reset_password'))
        else:
            flash('Respuesta incorrecta.', 'error')
    return render_template('security-question.html', pregunta=pregunta)


@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password():
    if not session.get('reset_allowed'):
        return redirect(url_for('forgot_password'))

    email = session.get('recovery_email')
    usuario = buscar_usuario_por_email(email)

    if request.method == 'POST':
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        if new_password != confirm_password:
            flash('Las contraseñas no coinciden.', 'error')
            return redirect(url_for('reset_password'))

        usuarios = cargar_usuarios()
        for u in usuarios:
            if u['email'] == email:
                u['password'] = generate_password_hash(new_password)
                break
        guardar_usuarios(usuarios)

        # Limpiar sesión temporal
        session.pop('recovery_email', None)
        session.pop('reset_allowed', None)

        flash('Contraseña restablecida correctamente. Inicia sesión.', 'success')
        return redirect(url_for('login'))

    return render_template('reset-password.html')


@app.route('/logout')
def logout():
    session.pop('usuario', None)
    flash('Sesión cerrada correctamente.', 'info')
    return redirect(url_for('login'))

#region CALENDARIO
@app.route('/calendar', methods=['GET', 'POST'])
def calendar():
    if 'usuario' not in session:
        flash('Debes iniciar sesión para acceder.', 'warning')
        return redirect(url_for('login'))
    
    # Cargar actualizaciones automáticas
    fechas_actualizaciones = []
    if os.path.exists(LOG_FILE):
        with open(LOG_FILE, 'r') as f:
            fechas_actualizaciones = json.load(f)
    eventos_actualizaciones = [{"title": "Se actualizó el archivo", "start": fecha} for fecha in fechas_actualizaciones]

    # Cargar eventos personalizados
    eventos_personalizados = []
    if os.path.exists(EVENTOS_FILE):
        with open(EVENTOS_FILE, 'r') as f:
            eventos_personalizados = json.load(f)

    todos_los_eventos = eventos_actualizaciones + eventos_personalizados

    hoy = date.today().isoformat()
    return render_template('calendar.html', eventos=todos_los_eventos, hoy=hoy)

def formatear_fecha_esp(fecha_iso):
    meses_rev = {v: k for k, v in MESES.items()}
    año, mes, dia = fecha_iso.split("-")
    mes_nombre = meses_rev.get(mes, "mes")
    return f"{int(dia)} de {mes_nombre} de {año}"

#region CREAR EVENTO
@app.route('/crear-evento', methods=['POST'])
def crear_evento():
    titulo = request.form['titulo']
    fecha = request.form['fecha']
    descripcion = request.form['descripcion']
    creador = session.get('usuario')

    nuevo_evento = {
        "title": titulo,
        "start": fecha,
        "description": descripcion,
        "creador": creador  # Nuevo campo
    }

    # Guardar en eventos.json
    if os.path.exists(EVENTOS_FILE):
        with open(EVENTOS_FILE, 'r') as f:
            eventos = json.load(f)
    else:
        eventos = []

    eventos.append(nuevo_evento)
    with open(EVENTOS_FILE, 'w') as f:
        json.dump(eventos, f, indent=4)

    # Enviar notificación por correo
    enviar_notificacion_email(titulo, fecha, descripcion)

    flash("Evento creado y notificación enviada.", "success")
    return redirect(url_for('calendar'))


#region --------------------- ENVIAR RECORDATORIOS DE EVENTOS -------------------------------
def enviar_recordatorios_eventos():
    with app.app_context():
        hoy = date.today().isoformat()
        if os.path.exists(EVENTOS_FILE):
            with open(EVENTOS_FILE, 'r') as f:
                eventos = json.load(f)
            for evento in eventos:
                if evento.get("start") == hoy and evento.get("creador"):
                    destinatario = evento["creador"]
                    try:
                        msg = Message(
                            subject=f"Recordatorio: {evento['title']}",
                            sender=app.config['MAIL_USERNAME'],
                            recipients=[destinatario]
                        )
                        msg.body = f"Hoy es el evento: {evento['title']}\n\nDescripción: {evento.get('description', '')}"
                        mail.send(msg)
                    except Exception as e:
                        print(f"Error enviando recordatorio: {e}")


# region ----------------------- ARCHIVO ----------------------------
@app.route('/archivo', methods=['GET', 'POST'])
def archivo():

    if 'usuario' not in session:
        flash('Debes iniciar sesión para acceder.', 'warning')
        return redirect(url_for('login'))
    
    # Inicializamos variables
    datos_estructurados = {}
    nombre_archivo_descarga = None
    

    if request.method == 'POST':
        file = request.files['archivo']
        if file:
            filename = secure_filename(file.filename)
            ruta_archivo = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(ruta_archivo)
            nombre_archivo_descarga = filename

            document = Document(ruta_archivo)
            parrafos = [p.text.strip().replace('\xa0', ' ') for p in document.paragraphs]
            contenido = "\n".join(parrafos)

            coincidencias = []
            rangos_ocupados = []

            sedes = ["Sede Cevaz La Limpia", "Sede Cevaz Digital", "Sede Cevaz Las Mercedes", "Sede Cevaz San Francisco", "General"]
            posiciones_sede = []
            posiciones_bloque = []

            def registrar_rango(start, end):
                rangos_ocupados.append((start, end))

            def esta_en_rango_ocupado(start, end):
                return any(start < r_end and end > r_start for r_start, r_end in rangos_ocupados)

            for i, parrafo in enumerate(parrafos):
                for sede in sedes:
                    if sede.lower() in parrafo.lower():
                        posiciones_sede.append((len("\n".join(parrafos[:i])), sede))
                if parrafo.isupper() and len(parrafo) > 5:
                    posiciones_bloque.append((len("\n".join(parrafos[:i])), parrafo.strip()))

            def obtener_sede(pos):
                return next((sede for p, sede in reversed(posiciones_sede) if pos >= p), "Sin sede")

            def obtener_bloque(pos):
                return next((bloque for p, bloque in reversed(posiciones_bloque) if pos >= p), "General")

            # --- Procesamiento de fechas ---
            for match in FECHA_RANGO_SIMPLE.finditer(contenido):
                dia_inicio, dia_fin, mes, año = match.groups()
                coincidencias.append({
                    "tipo": "simple",
                    "pos": match.start(),
                    "datos": (dia_inicio, dia_fin, mes, año),
                    "sede": obtener_sede(match.start()),
                    "bloque": obtener_bloque(match.start()),
                    "texto_original": match.group()
                })
                registrar_rango(match.start(), match.end())

            for match in FECHA_RANGO_DOBLE_MES.finditer(contenido):
                dia_inicio, mes_inicio, dia_fin, mes_fin, año = match.groups()
                coincidencias.append({
                    "tipo": "doble",
                    "pos": match.start(),
                    "datos": (dia_inicio, mes_inicio, dia_fin, mes_fin, año),
                    "sede": obtener_sede(match.start()),
                    "bloque": obtener_bloque(match.start()),
                    "texto_original": match.group()
                })
                registrar_rango(match.start(), match.end())

            for match in FECHA_UNICA.finditer(contenido):
                if not esta_en_rango_ocupado(match.start(), match.end()):
                    dia, mes, año = match.groups()
                    coincidencias.append({
                        "tipo": "unica",
                        "pos": match.start(),
                        "datos": (dia, mes, año),
                        "sede": obtener_sede(match.start()),
                        "bloque": obtener_bloque(match.start()),
                        "texto_original": match.group()
                    })
                    registrar_rango(match.start(), match.end())

            coincidencias.sort(key=lambda x: x["pos"])
            agrupado = {}

            ETIQUETAS_POR_BLOQUE = {
                "CALENDARIO DE INSCRIPCIONES: CLUB DE CONVERSACIÓN": ["Inscripción", "Inicio del club"]
            }

            ETIQUETAS_DEFAULT = ["Inscripción", "Regulares", "No regulares", "Inicio de clases"]

            for item in coincidencias:
                sede = item["sede"]
                bloque = item["bloque"]
                clave = (sede, bloque)
                if clave not in agrupado:
                    agrupado[clave] = []
                agrupado[clave].append(item)

            for (sede, bloque), items in agrupado.items():
                if sede not in datos_estructurados:
                    datos_estructurados[sede] = {}
                if bloque not in datos_estructurados[sede]:
                    datos_estructurados[sede][bloque] = []

                etiquetas_usar = ETIQUETAS_POR_BLOQUE.get(bloque.upper(), ETIQUETAS_DEFAULT)

                for i, item in enumerate(items):
                    if i >= len(etiquetas_usar):
                        break
                    etiqueta = etiquetas_usar[i]

                    if item["tipo"] == "simple":
                        dia_inicio, dia_fin, mes, año = item["datos"]
                        fecha_inicio = parse_fecha(dia_inicio, mes, año)
                        fecha_fin = parse_fecha(dia_fin, mes, año)
                    elif item["tipo"] == "doble":
                        dia_inicio, mes_inicio, dia_fin, mes_fin, año = item["datos"]
                        fecha_inicio = parse_fecha(dia_inicio, mes_inicio, año)
                        fecha_fin = parse_fecha(dia_fin, mes_fin, año)
                    else:
                        dia, mes, año = item["datos"]
                        fecha_inicio = parse_fecha(dia, mes, año)
                        fecha_fin = None

                    datos_estructurados[sede][bloque].append({
                        "etiqueta": etiqueta,
                        "fecha_inicio": fecha_inicio,
                        "fecha_fin": fecha_fin,
                        "original": item["texto_original"]
                    })

            # --- Procesamiento de precios ---
            precios_agrupados = {}

            for i, parrafo in enumerate(parrafos):
                pos = len("\n".join(parrafos[:i]))
                sede_actual = obtener_sede(pos)
                bloque_actual = obtener_bloque(pos)
                clave = (sede_actual, bloque_actual)

                precios = PRECIO_REGEX.findall(parrafo)
                impuesto = IMPUESTO_REGEX.search(parrafo)
                precio_total = PRECIO_TOTAL_REGEX.search(parrafo)
                
                if precios:
                    if clave not in precios_agrupados:
                        precios_agrupados[clave] = []

                    precio_base = precios[0]
                    porcentaje_impuesto = impuesto.group(1) if impuesto else None

                    # Cálculo automático del total si impuesto está presente
                    if porcentaje_impuesto:
                        total = round(float(precio_base) * (1 + float(porcentaje_impuesto) / 100), 2)
                    else:
                        total = None  # No se puede calcular sin impuesto

                    precios_agrupados[clave].append({
                        "es_precio": True,
                        "etiqueta": "Precio del Servicio",
                        "precio_base": precio_base,
                        "porcentaje_impuesto": porcentaje_impuesto,
                        "precio_total": f"{total:.2f}" if total is not None else None,
                        "original": parrafo
                    })


            # Insertar en la estructura principal
            for clave, precios in precios_agrupados.items():
                sede, bloque = clave
                if sede not in datos_estructurados:
                    datos_estructurados[sede] = {}
                if bloque not in datos_estructurados[sede]:
                    datos_estructurados[sede][bloque] = []

                datos_estructurados[sede][bloque].extend(precios)
                
    hoy = date.today().isoformat()
    return render_template('archivo.html', datos=datos_estructurados, archivo_subido=nombre_archivo_descarga, hoy=hoy)


# region ACTUALIZAR FECHAS Y PRECIOS
@app.route('/actualizar-fechas', methods=['POST'])
def actualizar_fechas():
    archivo_original = request.form['archivo_original']
    ruta_archivo = os.path.join(app.config['UPLOAD_FOLDER'], archivo_original)
    document = Document(ruta_archivo)

    nuevas_fechas = {}
    nuevos_precios = {}
    
    for key in request.form:
        if key.startswith("fecha_inicio-"):
            partes = key.split("-")
            sede, bloque, idx = partes[1], partes[2], partes[3]
            fecha_inicio = request.form[key]
            fecha_fin = request.form.get(f"fecha_fin-{sede}-{bloque}-{idx}")
            etiqueta = request.form[f"meta-{sede}-{bloque}-{idx}-etiqueta"]
            texto_original = request.form[f"original-{sede}-{bloque}-{idx}"]
            nuevas_fechas[(sede, bloque, etiqueta)] = (fecha_inicio, fecha_fin, texto_original)
            
    for key in request.form:
        if key.startswith("precio_base-"):
            partes = key.split("-")
            sede, bloque, idx = partes[1], partes[2], partes[3]
            precio_base = request.form[key]
            porcentaje_impuesto = request.form.get(f"porcentaje_impuesto-{sede}-{bloque}-{idx}")
            precio_total = None  # <-- Inicializa aquí
            # Recalcula el total si base e impuesto están presentes
            try:
                if precio_base and porcentaje_impuesto:
                    total = round(float(precio_base) * (1 + float(porcentaje_impuesto) / 100), 2)
                    precio_total = f"{total:.2f}"
            except Exception as e:
                print(f"Error calculando total: {e}")
                precio_total = None

            texto_original = request.form.get(f"original-precio-{sede}-{bloque}-{idx}")
            nuevos_precios[(sede, bloque, idx)] = (precio_base, porcentaje_impuesto, precio_total, texto_original)
     
    for p in document.paragraphs:
        for (sede, bloque, etiqueta), (nueva_ini, nueva_fin, texto_original) in nuevas_fechas.items():
            if texto_original in p.text:
                nueva_fecha = ""
                if nueva_ini and nueva_fin:
                    nueva_fecha = f"{formatear_fecha_esp(nueva_ini)} al {formatear_fecha_esp(nueva_fin)}"
                elif nueva_ini:
                    nueva_fecha = formatear_fecha_esp(nueva_ini)
                p.text = p.text.replace(texto_original, nueva_fecha)
    
    for p in document.paragraphs:
        for (sede, bloque, idx), (precio_base, porcentaje_impuesto, precio_total, texto_original) in nuevos_precios.items():
            if texto_original and texto_original in p.text:
                nuevo_texto = p.text
                # Reemplaza solo el primer precio base
                if precio_base:
                    nuevo_texto = re.sub(r"\$[0-9]+(?:\.[0-9]{1,2})?", f"${precio_base}", nuevo_texto, count=1)
                # Reemplaza solo el porcentaje de impuesto
                if porcentaje_impuesto:
                    nuevo_texto = re.sub(r"(\d{1,2}(?:\.\d{1,2})?)\s*%(\s*de)?\s*(impuesto[s]?|IGTF)", f"{porcentaje_impuesto}% de \\3", nuevo_texto)
                # Reemplaza solo el monto total
                if precio_total:
                    nuevo_texto = re.sub(
                        r"(monto a pagar (?:sería|es|seria|seran|serán)[^\$]*\$)([0-9]+(?:\.[0-9]{1,2})?)",
                        lambda m: f"{m.group(1)}{precio_total}",
                        nuevo_texto
                    )
                p.text = nuevo_texto

    fecha_hoy = datetime.today().strftime('%d-%m-%Y_%H-%M-%S')
    nuevo_nombre = f"VARIABLE AL {fecha_hoy}.docx"
    ruta_transaccion = os.path.join(app.config['TRANSACCIONES_FOLDER'], nuevo_nombre)
    document.save(ruta_transaccion)
    registrar_actualizacion()
    
    
    # Enviar el archivo actualizado por correo al usuario logueado
    email_usuario = session.get('usuario')
    if email_usuario:
        enviar_archivo_actualizado_email(email_usuario, ruta_transaccion, nuevo_nombre)


    return send_file(ruta_transaccion, as_attachment=True, download_name=nuevo_nombre)



#region TRANSACCIONES
@app.route('/transacciones')
def transacciones():
    if 'usuario' not in session:
        flash('Debes iniciar sesión para acceder.', 'warning')
        return redirect(url_for('login'))
    
    archivos = sorted(os.listdir(app.config['TRANSACCIONES_FOLDER']), reverse=True)
    return render_template('transacciones.html', archivos=archivos)


#region DESCARGAR TRANSACCIONES
@app.route('/descargar-transaccion/<nombre_archivo>')
def descargar_transaccion(nombre_archivo):
    ruta_archivo = os.path.join(app.config['TRANSACCIONES_FOLDER'], nombre_archivo)
    return send_file(ruta_archivo, as_attachment=True, download_name=nombre_archivo)


#region CONFIGURACIONES
@app.route('/settings', methods=['GET', 'POST'])
def settings():
    if 'usuario' not in session:
        flash('Debes iniciar sesión para acceder.', 'warning')
        return redirect(url_for('login'))

    usuario = buscar_usuario_por_email(session['usuario'])

    # Lista de preguntas predeterminadas
    preguntas_predeterminadas = [
        "¿Cuál es el nombre de tu primera mascota?",
        "¿En qué ciudad naciste?",
        "¿Cuál es tu comida favorita?"
    ]

    if request.method == 'POST':
        action = request.form.get('action')
        usuarios = cargar_usuarios()
        user = next((u for u in usuarios if u['email'] == usuario['email']), None)

        if action == 'change_email':
            new_email = request.form['new_email']
            if buscar_usuario_por_email(new_email):
                flash('Ese correo ya está registrado.', 'error')
            else:
                user['email'] = new_email
                session['usuario'] = new_email
                guardar_usuarios(usuarios)
                flash('Correo actualizado correctamente.', 'success')

        elif action == 'change_password':
            current_password = request.form['current_password']
            new_password = request.form['new_password']
            confirm_password = request.form['confirm_password']
            if not check_password_hash(user['password'], current_password):
                flash('Contraseña actual incorrecta.', 'error')
            elif new_password != confirm_password:
                flash('Las contraseñas nuevas no coinciden.', 'error')
            else:
                user['password'] = generate_password_hash(new_password)
                guardar_usuarios(usuarios)
                flash('Contraseña actualizada correctamente.', 'success')

        elif action == 'change_recovery':
            pregunta = request.form['pregunta_recuperacion']
            respuesta = request.form['respuesta_recuperacion']
            user['pregunta_recuperacion'] = pregunta
            user['respuesta_recuperacion'] = respuesta
            guardar_usuarios(usuarios)
            flash('Pregunta de recuperación actualizada.', 'success')

        return redirect(url_for('settings'))

    # Pasa la pregunta y respuesta actual del usuario
    return render_template(
        'settings.html',
        preguntas_predeterminadas=preguntas_predeterminadas,
        pregunta_actual=usuario.get('pregunta_recuperacion', ''),
        respuesta_actual=usuario.get('respuesta_recuperacion', '')
    )

#region HELP
@app.route('/help')
def help():
    if 'usuario' not in session:
        flash('Debes iniciar sesión para acceder.', 'warning')
        return redirect(url_for('login'))
    
    return render_template('help.html')

#region ENVIAR ARCHIVO ACTUALIZADO POR CORREO
def enviar_archivo_actualizado_email(destinatario, ruta_archivo, nombre_archivo):
    try:
        msg = Message(
            subject="Archivo maestro actualizado",
            sender=app.config['MAIL_USERNAME'],
            recipients=[destinatario]
        )
        msg.body = "El archivo maestro ha sido actualizado, por favor actualizar el script en chatbot by magic school."
        with open(ruta_archivo, "rb") as fp:
            msg.attach(nombre_archivo, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fp.read())
        mail.send(msg)
    except Exception as e:
        print(f"Error al enviar el archivo actualizado por correo: {str(e)}")
        flash("Error al enviar el archivo actualizado por correo.", "error")

#region ENVIAR NOTIFICACIÓN POR CORREO
def enviar_notificacion_email(titulo, fecha, descripcion):
    try:
        # Obtener el correo del usuario actualmente logueado
        email_usuario = session.get('usuario')

        if not email_usuario:
            flash("No se pudo enviar el correo porque no hay usuario logueado.", "error")
            return

        msg = Message(
            subject=f"Nuevo evento: {titulo}",
            sender=app.config['MAIL_USERNAME'],
            recipients=[email_usuario]  # se envía al usuario actual
        )
        msg.body = f"Evento programado para el {fecha}:\n\n{descripcion}"
        mail.send(msg)
    except Exception as e:
        print(f"Error al enviar correo: {str(e)}")
        flash("Error al enviar notificación por correo.", "error")

#region DESCARGAR ARCHIVO
@app.route('/descargar/<nombre_archivo>')
def descargar_archivo(nombre_archivo):
    fecha_hoy = datetime.today().strftime('%d-%m-%Y_%H-%M-%S')
    nuevo_nombre = f"VARIABLE AL {fecha_hoy}.docx"
    ruta_archivo = os.path.join(app.config['UPLOAD_FOLDER'], nombre_archivo)
    return send_file(ruta_archivo, as_attachment=True, download_name=nuevo_nombre)


# region ------------------------ INICIO ------------------------
if __name__ == '__main__':
    scheduler = BackgroundScheduler()
    scheduler.add_job(enviar_recordatorios_eventos, 'cron', hour=8, minute=0)  # Enviar recordatorios todos los días a las 8:00 AM
    scheduler.start()
    app.run(debug=True)